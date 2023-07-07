# 读取word数据
from docx import Document

# 1.创建doc对象
doc = Document(r'data/占勇辉的离职证明.docx')

# 2.获取doc中所有的段落，结果是一个列表，其中的元素是段落对象
print(doc.paragraphs)

# 2.遍历段落,获取段落的文本组成
for i,p in enumerate(doc.paragraphs):
    print(i, p.text)
    # 获取每个段落的组件
    print(p.runs)
    for run in p.runs:
        print(run.text)     # 获取每个组件的文本内容
