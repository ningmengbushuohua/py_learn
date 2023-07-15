from docx import Document
from docx.shared import Cm
# pip install python-docx -i https://pypi.tuna.tsinghua.edu.cn/simple/

# 写入word
# 注意:操作word,需要安装python-docx库，但是在代码中使用的时候导入docx

# 1. 创建doc 文档对象
doc =Document()
# 2.添加标题
doc.add_heading('在python中操作word文档',2)

# 3.添加段落
p = doc.add_paragraph('python是一种面向对象的语言，')
# 注意：在word中，一个段落可能有很多组件组成
# 给段落添加组件
run1 = p.add_run('python是可以跨平台的,')
run2 = p.add_run('python可以做数据分析,')
run3 = p.add_run('python可以做web开发,等')
run1.bold =True     # 设置加粗
run2.underline = True   # 设置下划线


# 4.添加列表
# 有序列表
doc.add_paragraph('运营部',style='List Number')
doc.add_paragraph('行政部',style='List Number')
doc.add_paragraph('市场部',style='List Number')
doc.add_paragraph('财务部',style='List Number')

# 无序列表
doc.add_paragraph('运营部',style='List Bullet')
doc.add_paragraph('行政部',style='List Bullet')
doc.add_paragraph('市场部',style='List Bullet')
doc.add_paragraph('财务部',style='List Bullet')

# 5.添加图片
doc.add_picture(r'data/pic.png',width=Cm(3))

# 保存文档
doc.save(r'data/python中操作word文档.docx')

#
